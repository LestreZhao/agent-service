import io
import uuid
from datetime import datetime, timedelta
from typing import Dict, Any, Optional
from pathlib import Path
import logging
import traceback

import PyPDF2
from docx import Document
from minio.error import S3Error

# 导入日志配置
from src.utils.logger_config import setup_logging
from src.config.minio import get_minio_client, MINIO_BUCKET_NAME, ensure_bucket_exists

# 配置日志记录
logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)


class DocumentParser:
    """文档解析器类"""
    
    def __init__(self):
        self.minio_client = get_minio_client()
        ensure_bucket_exists()
    
    def extract_pdf_content(self, file_content: bytes) -> str:
        """提取PDF文档内容"""
        try:
            logger.debug(f"开始解析PDF，文件大小: {len(file_content)} 字节")
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            page_count = len(pdf_reader.pages)
            logger.debug(f"PDF总页数: {page_count}")
            
            text = ""
            for i, page in enumerate(pdf_reader.pages):
                logger.debug(f"正在解析第 {i+1} 页")
                page_text = page.extract_text()
                text += page_text + "\n"
                logger.debug(f"第 {i+1} 页提取了 {len(page_text)} 个字符")
            
            result = text.strip()
            logger.debug(f"PDF解析完成，总共提取了 {len(result)} 个字符")
            return result
        except Exception as e:
            logger.error(f"PDF解析错误: {str(e)}")
            logger.error(f"详细错误信息: {traceback.format_exc()}")
            raise ValueError(f"PDF解析错误: {str(e)}")
    
    def extract_docx_content(self, file_content: bytes) -> str:
        """提取Word文档内容（包括表格）"""
        try:
            logger.debug(f"开始解析Word文档，文件大小: {len(file_content)} 字节")
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            logger.debug(f"Word文档段落数量: {paragraph_count}, 表格数量: {table_count}")
            
            text = ""
            
            # 提取段落内容
            for i, paragraph in enumerate(doc.paragraphs):
                logger.debug(f"正在解析第 {i+1} 个段落")
                paragraph_text = paragraph.text.strip()
                if paragraph_text:  # 只添加非空段落
                    text += paragraph_text + "\n"
                    logger.debug(f"第 {i+1} 个段落提取了 {len(paragraph_text)} 个字符")
            
            # 提取表格内容
            for i, table in enumerate(doc.tables):
                logger.debug(f"正在解析第 {i+1} 个表格")
                text += f"\n--- 表格 {i+1} ---\n"
                
                for row_idx, row in enumerate(table.rows):
                    row_text = []
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        row_text.append(cell_text)
                        logger.debug(f"表格 {i+1} 第 {row_idx+1} 行第 {cell_idx+1} 列: {len(cell_text)} 个字符")
                    
                    # 使用制表符分隔单元格内容
                    text += "\t".join(row_text) + "\n"
                
                text += f"--- 表格 {i+1} 结束 ---\n\n"
                logger.debug(f"第 {i+1} 个表格解析完成")
            
            result = text.strip()
            logger.debug(f"Word文档解析完成，总共提取了 {len(result)} 个字符")
            logger.info(f"解析结果: {paragraph_count} 个段落, {table_count} 个表格")
            return result
        except Exception as e:
            logger.error(f"Word文档解析错误: {str(e)}")
            logger.error(f"详细错误信息: {traceback.format_exc()}")
            raise ValueError(f"Word文档解析错误: {str(e)}")
    
    def extract_docx_content_advanced(self, file_content: bytes) -> str:
        """高级Word文档内容提取（更好的表格处理）"""
        try:
            logger.debug(f"开始高级解析Word文档，文件大小: {len(file_content)} 字节")
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text_content = []
            
            # 获取文档中所有元素（按顺序）
            elements = []
            
            # 先处理所有段落和表格的相对位置
            for element in doc.element.body:
                if element.tag.endswith('p'):  # 段落
                    # 找到对应的段落对象
                    for para in doc.paragraphs:
                        if para._element == element:
                            elements.append(('paragraph', para))
                            break
                elif element.tag.endswith('tbl'):  # 表格
                    # 找到对应的表格对象
                    for table in doc.tables:
                        if table._element == element:
                            elements.append(('table', table))
                            break
            
            logger.debug(f"找到 {len(elements)} 个文档元素")
            
            # 按顺序处理元素
            for i, (element_type, element) in enumerate(elements):
                if element_type == 'paragraph':
                    para_text = element.text.strip()
                    if para_text:
                        text_content.append(para_text)
                        logger.debug(f"段落 {i+1}: {len(para_text)} 个字符")
                
                elif element_type == 'table':
                    table_text = self._format_table_content(element, i+1)
                    if table_text:
                        text_content.append(table_text)
                        logger.debug(f"表格 {i+1}: {len(table_text)} 个字符")
            
            result = "\n\n".join(text_content)
            logger.debug(f"高级Word文档解析完成，总共提取了 {len(result)} 个字符")
            return result
            
        except Exception as e:
            logger.warning(f"高级解析失败，回退到基础解析: {str(e)}")
            # 如果高级解析失败，回退到基础解析
            return self.extract_docx_content(file_content)
    
    def _format_table_content(self, table, table_num: int) -> str:
        """格式化表格内容"""
        try:
            table_lines = []
            table_lines.append(f"【表格 {table_num}】")
            
            # 检查是否有表头
            if len(table.rows) > 0:
                first_row = table.rows[0]
                headers = []
                for cell in first_row.cells:
                    headers.append(cell.text.strip())
                
                # 如果第一行看起来像表头（内容较短且不包含大量数字），则特殊处理
                is_header_row = self._is_likely_header_row(headers)
                
                if is_header_row:
                    # 添加表头
                    table_lines.append("表头: " + " | ".join(headers))
                    table_lines.append("-" * 50)
                    start_row = 1
                else:
                    start_row = 0
                
                # 处理数据行
                for row_idx in range(start_row, len(table.rows)):
                    row = table.rows[row_idx]
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip().replace('\n', ' ').replace('\t', ' ')
                        row_data.append(cell_text)
                    
                    if any(row_data):  # 只添加非空行
                        if is_header_row and len(headers) == len(row_data):
                            # 如果有表头，使用键值对格式
                            row_content = []
                            for header, data in zip(headers, row_data):
                                if data:  # 只显示非空数据
                                    row_content.append(f"{header}: {data}")
                            if row_content:
                                table_lines.append(" | ".join(row_content))
                        else:
                            # 否则使用简单的制表符分隔
                            table_lines.append(" | ".join(row_data))
            
            return "\n".join(table_lines)
            
        except Exception as e:
            logger.error(f"表格格式化失败: {str(e)}")
            # 回退到简单格式
            simple_content = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                simple_content.append("\t".join(row_text))
            return f"【表格 {table_num}】\n" + "\n".join(simple_content)
    
    def _is_likely_header_row(self, row_data: list) -> bool:
        """判断是否为表头行"""
        if not row_data:
            return False
        
        # 检查条件：
        # 1. 大部分单元格都有内容
        # 2. 内容相对较短
        # 3. 不全是数字
        non_empty_count = sum(1 for cell in row_data if cell.strip())
        if non_empty_count < len(row_data) * 0.5:
            return False
        
        # 检查是否主要是数字
        numeric_count = 0
        for cell in row_data:
            cell = cell.strip()
            if cell and cell.replace('.', '').replace('-', '').replace(',', '').isdigit():
                numeric_count += 1
        
        # 如果超过一半是数字，可能不是表头
        if numeric_count > len(row_data) * 0.5:
            return False
        
        # 检查长度
        avg_length = sum(len(cell.strip()) for cell in row_data) / len(row_data)
        if avg_length > 50:  # 如果平均长度太长，可能不是表头
            return False
        
        return True
    
    def parse_document(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """解析文档并提取内容"""
        logger.debug(f"开始解析文档: {filename}")
        file_ext = Path(filename).suffix.lower()
        logger.debug(f"检测到文件类型: {file_ext}")
        
        try:
            if file_ext == '.pdf':
                logger.debug("使用PDF解析器")
                content = self.extract_pdf_content(file_content)
            elif file_ext in ['.docx', '.doc']:
                logger.debug("使用Word文档解析器")
                content = self.extract_docx_content_advanced(file_content)
            else:
                logger.error(f"不支持的文件类型: {file_ext}")
                raise ValueError(f"不支持的文件类型: {file_ext}")
            
            logger.debug(f"解析完成，提取内容长度: {len(content)}")
            
            return {
                "filename": filename,
                "file_type": file_ext,
                "content": content,
                "content_length": len(content),
                "parsed_at": datetime.now().isoformat()
            }
        except Exception as e:
            logger.error(f"文档解析失败: {str(e)}")
            logger.error(f"详细错误信息: {traceback.format_exc()}")
            raise
    
    def upload_file(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """上传文件到MinIO并解析"""
        try:
            # 生成唯一的文件ID
            file_id = str(uuid.uuid4())
            object_name = f"{file_id}_{filename}"
            
            # 上传到MinIO
            file_stream = io.BytesIO(file_content)
            self.minio_client.put_object(
                MINIO_BUCKET_NAME,
                object_name,
                file_stream,
                length=len(file_content),
                content_type=self._get_content_type(filename)
            )
            
            # 解析文档内容
            document_info = self.parse_document(file_content, filename)
            document_info.update({
                "file_id": file_id,
                "object_name": object_name,
                "file_size": len(file_content),
                "uploaded_at": datetime.now().isoformat()
            })
            
            return document_info
            
        except S3Error as e:
            raise ValueError(f"文件上传失败: {str(e)}")
        except Exception as e:
            raise ValueError(f"文档处理失败: {str(e)}")
    
    def get_file_info(self, file_id: str) -> Optional[Dict[str, Any]]:
        """根据文件ID获取文件信息"""
        logger.info(f"开始获取文件信息，文件ID: {file_id}")
        
        try:
            # 检查MinIO连接
            logger.debug(f"检查MinIO连接，存储桶: {MINIO_BUCKET_NAME}")
            
            # 查找以file_id开头的对象
            logger.debug(f"在存储桶中搜索文件，前缀: {file_id}")
            objects = self.minio_client.list_objects(MINIO_BUCKET_NAME, prefix=file_id)
            
            objects_list = list(objects)
            logger.debug(f"找到 {len(objects_list)} 个匹配的对象")
            
            for obj in objects_list:
                logger.debug(f"检查对象: {obj.object_name}")
                if obj.object_name.startswith(file_id):
                    logger.info(f"找到匹配的文件: {obj.object_name}")
                    
                    try:
                        # 获取文件内容
                        logger.debug("开始下载文件内容")
                        response = self.minio_client.get_object(MINIO_BUCKET_NAME, obj.object_name)
                        file_content = response.read()
                        response.close()
                        logger.debug(f"下载完成，文件大小: {len(file_content)} 字节")
                        
                        # 解析原始文件名
                        original_filename = obj.object_name[len(file_id) + 1:]  # 去掉file_id和下划线
                        logger.debug(f"解析得到原始文件名: {original_filename}")
                        
                        # 解析文档内容
                        logger.debug("开始解析文档内容")
                        document_info = self.parse_document(file_content, original_filename)
                        
                        document_info.update({
                            "file_id": file_id,
                            "object_name": obj.object_name,
                            "file_size": obj.size,
                            "last_modified": obj.last_modified.isoformat() if obj.last_modified else None
                        })
                        
                        logger.info(f"文档解析成功，内容长度: {len(document_info.get('content', ''))}")
                        return document_info
                        
                    except Exception as content_error:
                        logger.error(f"处理文件内容时发生错误: {str(content_error)}")
                        logger.error(f"详细错误信息: {traceback.format_exc()}")
                        raise ValueError(f"处理文件内容失败: {str(content_error)}")
            
            logger.warning(f"未找到匹配的文件，文件ID: {file_id}")
            return None
            
        except S3Error as e:
            logger.error(f"MinIO S3错误: {str(e)}")
            logger.error(f"详细错误信息: {traceback.format_exc()}")
            raise ValueError(f"获取文件信息失败 (S3错误): {str(e)}")
        except Exception as e:
            logger.error(f"未知错误: {str(e)}")
            logger.error(f"详细错误信息: {traceback.format_exc()}")
            raise ValueError(f"文档处理失败: {str(e)}")
    
    def download_file(self, file_id: str) -> Optional[bytes]:
        """下载文件内容"""
        try:
            objects = self.minio_client.list_objects(MINIO_BUCKET_NAME, prefix=file_id)
            
            for obj in objects:
                if obj.object_name.startswith(file_id):
                    response = self.minio_client.get_object(MINIO_BUCKET_NAME, obj.object_name)
                    file_content = response.read()
                    response.close()
                    return file_content
            
            return None
            
        except S3Error as e:
            raise ValueError(f"下载文件失败: {str(e)}")
    
    def delete_file(self, file_id: str) -> bool:
        """删除文件"""
        try:
            objects = self.minio_client.list_objects(MINIO_BUCKET_NAME, prefix=file_id)
            
            for obj in objects:
                if obj.object_name.startswith(file_id):
                    self.minio_client.remove_object(MINIO_BUCKET_NAME, obj.object_name)
                    return True
            
            return False
            
        except S3Error as e:
            raise ValueError(f"删除文件失败: {str(e)}")
    
    def _get_content_type(self, filename: str) -> str:
        """根据文件名获取内容类型"""
        ext = Path(filename).suffix.lower()
        content_types = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword'
        }
        return content_types.get(ext, 'application/octet-stream')


# 创建全局实例
document_parser = DocumentParser() 